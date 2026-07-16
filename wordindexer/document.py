"""
Document reader.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph as DocxParagraph

from wordindexer.exceptions import DocumentError
from wordindexer.models import Book, Heading, Paragraph


@dataclass(slots=True)
class DocumentInfo:
    title: str
    author: str
    subject: str
    paragraphs: int
    tables: int
    images: int


class DocumentReader:
    """Read a DOCX document into WordIndexer domain objects."""

    def __init__(self, filename: str | Path):
        self.path = Path(filename)

        if not self.path.exists():
            raise DocumentError(f"Document not found: {self.path}")

        self.doc = Document(self.path)

    def inspect(self) -> DocumentInfo:
        props = self.doc.core_properties
        image_count = 0

        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1

        return DocumentInfo(
            title=props.title or "",
            author=props.author or "",
            subject=props.subject or "",
            paragraphs=len(self.doc.paragraphs),
            tables=len(self.doc.tables),
            images=image_count,
        )

    def _iter_body_blocks(self) -> Iterator[DocxParagraph | Table]:
        """Yield top-level paragraphs and tables in document order."""
        body = self.doc.element.body

        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield DocxParagraph(child, self.doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, self.doc)

    @staticmethod
    def _heading_level(style: str) -> int | None:
        if not style.startswith("Heading"):
            return None

        try:
            return int(style.split()[-1])
        except Exception:
            return 1

    def _append_paragraph(
        self,
        book: Book,
        paragraph: DocxParagraph,
        index: int,
        *,
        story: str = "body",
        table_index: int | None = None,
        row_index: int | None = None,
        cell_index: int | None = None,
        cell_paragraph_index: int | None = None,
    ) -> int:
        style = paragraph.style.name if paragraph.style else ""

        book.paragraphs.append(
            Paragraph(
                index=index,
                text=paragraph.text,
                style=style,
                story=story,
                table_index=table_index,
                row_index=row_index,
                cell_index=cell_index,
                cell_paragraph_index=cell_paragraph_index,
            )
        )
        book.paragraph_targets[index] = paragraph

        level = self._heading_level(style)
        if level is not None:
            book.headings.append(
                Heading(
                    level=level,
                    text=paragraph.text,
                    paragraph_index=index,
                )
            )

        return index + 1

    def _append_table(
        self,
        book: Book,
        table: Table,
        index: int,
        table_index: int,
    ) -> int:
        """Append unique cell paragraphs from one top-level table."""
        seen_cells: set[int] = set()

        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_identity = id(cell._tc)

                if cell_identity in seen_cells:
                    continue

                seen_cells.add(cell_identity)

                for cell_paragraph_index, paragraph in enumerate(
                    cell.paragraphs
                ):
                    index = self._append_paragraph(
                        book,
                        paragraph,
                        index,
                        story="table",
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=cell_index,
                        cell_paragraph_index=cell_paragraph_index,
                    )

        return index

    def load_book(self, include_tables: bool = False) -> Book:
        """
        Load document paragraphs into a Book in document order.

        Top-level body paragraphs are always loaded. Table-cell paragraphs
        are loaded only when ``include_tables`` is true. When included, table
        paragraphs receive positions in the same document-order stream as
        body paragraphs, so TOC/body boundaries remain correct.
        """
        props = self.doc.core_properties
        book = Book(
            title=props.title or "",
            author=props.author or "",
        )

        next_index = 0
        table_index = 0

        for block in self._iter_body_blocks():
            if isinstance(block, DocxParagraph):
                next_index = self._append_paragraph(
                    book,
                    block,
                    next_index,
                )
            elif include_tables:
                next_index = self._append_table(
                    book,
                    block,
                    next_index,
                    table_index,
                )
                table_index += 1

        return book
