"""
Append a human-readable glossary section to a Word document.
"""

from __future__ import annotations

from docx import Document
from docx.text.paragraph import Paragraph

from wordindexer.glossary import GlossaryReport


class GlossaryWriter:
    """Write glossary entries as formatted Word paragraphs."""

    def append(
        self,
        document: Document,
        report: GlossaryReport,
        heading: str = "Glossary",
    ) -> Paragraph:
        """Append a glossary heading and defined entries to a document."""
        heading_paragraph = document.add_paragraph(
            heading,
            style="Heading 1",
        )
        heading_paragraph.paragraph_format.page_break_before = True

        for entry in report.entries:
            paragraph = document.add_paragraph()
            term_run = paragraph.add_run(entry.term)
            term_run.bold = True

            if entry.category:
                paragraph.add_run(f" ({entry.category})")

            paragraph.add_run(f": {entry.definition}")

        return heading_paragraph
