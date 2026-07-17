from docx import Document

from wordindexer.models import DictionaryEntry
from wordindexer.reports import ReportBuilder


def test_report_can_include_table_cell_occurrences(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("Body text")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("PowerShell in a table")
    document.save(source)

    report = ReportBuilder().build(
        source,
        [DictionaryEntry(term="PowerShell", index_as="PowerShell")],
        include_tables=True,
    )

    assert report.total_occurrences == 1
    assert report.term_counts["PowerShell"] == 1
