from docx import Document

from wordindexer.xmlwriter import XMLWriter


def test_insert_field_preserves_visible_text_and_run_formatting():
    document = Document()
    paragraph = document.add_paragraph()

    before = paragraph.add_run("Before ")
    before.bold = True
    paragraph.add_run("PowerShell after")

    XMLWriter().insert_field(
        paragraph=paragraph,
        run_index=1,
        offset=0,
        field_code='XE "PowerShell"',
    )

    assert paragraph.text == "Before PowerShell after"

    visible_runs = [run for run in paragraph.runs if run.text]
    assert visible_runs[0].text == "Before "
    assert visible_runs[0].bold is True
    assert visible_runs[-1].text == "PowerShell after"

    xml = paragraph._p.xml

    assert 'w:fldCharType="begin"' in xml
    assert 'w:fldCharType="end"' in xml
    assert 'XE "PowerShell"' in xml

    field_instructions = paragraph._p.xpath(".//w:instrText")

    assert len(field_instructions) == 1
    assert field_instructions[0].text == ' XE "PowerShell" '


def test_insert_field_supports_word_style_instruction_parts():
    document = Document()
    paragraph = document.add_paragraph("PowerShell")

    XMLWriter().insert_field(
        paragraph=paragraph,
        run_index=0,
        offset=len("PowerShell"),
        field_code='XE "PowerShell"',
        instruction_parts=[
            ' XE "',
            "PowerShell",
            '" ',
        ],
    )

    assert paragraph.text == "PowerShell"
    assert [
        element.text
        for element in paragraph._p.xpath(".//w:instrText")
    ] == [
        ' XE "',
        "PowerShell",
        '" ',
    ]
