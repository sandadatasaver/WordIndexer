from docx import Document

from wordindexer.models import DictionaryEntry
from wordindexer.workflow import DocumentWorkflow


def test_document_workflow_can_generate_glossary_docx(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"

    document = Document()
    document.add_paragraph("PowerShell appears here.")
    document.save(source)

    result = DocumentWorkflow().run(
        source,
        [
            DictionaryEntry(
                term="PowerShell",
                definition="A shell.",
            )
        ],
        output,
        generate_index=False,
        generate_glossary=True,
    )

    assert result.output_path == output
    assert result.glossary_entries == 1

    generated = Document(output)
    assert "Glossary" in [paragraph.text for paragraph in generated.paragraphs]
    assert any(
        paragraph.text == "PowerShell: A shell."
        for paragraph in generated.paragraphs
    )
