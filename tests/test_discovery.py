import json

from docx import Document

from wordindexer.discovery import TermDiscovery
from wordindexer.models import DictionaryEntry


def test_term_discovery_returns_reviewable_candidates(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph(
        "PowerShell uses Get-ChildItem. PowerShell is useful."
    )
    document.add_paragraph("PowerShell and Get-ChildItem appear again.")
    document.save(source)

    report = TermDiscovery(minimum_occurrences=2).discover(source)

    terms = {candidate.term: candidate for candidate in report.candidates}

    assert "PowerShell" in terms
    assert "Get-ChildItem" in terms
    assert terms["PowerShell"].occurrences == 3
    assert terms["Get-ChildItem"].occurrences == 2
    assert terms["PowerShell"].suggested_entry()["index_as"] == "PowerShell"
    assert report.body_start == 0


def test_discovery_excludes_path_fragments(tmp_path):
    source = tmp_path / "paths.docx"
    document = Document()
    document.add_paragraph(
        r"Use C:\PowerShell-Practice\Reports."
    )
    document.add_paragraph(
        "PowerShell-Practice is a path used in examples."
    )
    document.save(source)

    report = TermDiscovery(minimum_occurrences=2).discover(source)

    assert all(
        candidate.term != "PowerShell-Practice"
        for candidate in report.candidates
    )


def test_discovery_excludes_existing_dictionary_terms_and_writes_json(
    tmp_path,
):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("PowerShell and New-Term appear twice.")
    document.add_paragraph("New-Term appears again.")
    document.save(source)

    existing = [
        DictionaryEntry(term="PowerShell", index_as="PowerShell"),
    ]

    report = TermDiscovery().discover(source, existing)
    output = report.write_json(tmp_path / "discovery.json")
    payload = json.loads(output.read_text(encoding="utf-8"))

    assert [candidate.term for candidate in report.candidates] == [
        "New-Term",
    ]
    assert payload["candidates"][0]["suggested_entry"]["term"] == "New-Term"
