from docx import Document

from wordindexer.backmatter import BackMatterWriter
from wordindexer.glossary import GlossaryEntry, GlossaryReport


def test_backmatter_writer_replaces_existing_index_and_orders_glossary(
    tmp_path,
):
    document = Document()
    document.add_paragraph("Recommended Resources", style="Heading 1")
    document.add_paragraph("Index", style="Heading 1")
    document.add_paragraph("Old index content")
    document.add_paragraph("About the Author", style="Heading 1")
    document.add_paragraph("Author biography")
    document.add_paragraph("Contact", style="Heading 1")
    document.add_paragraph("Contact information")

    report = GlossaryReport(
        input_path="source.docx",
        entries=[
            GlossaryEntry(
                term="PowerShell",
                definition="A shell.",
                category="Technology",
                aliases=[],
                occurrences=2,
                found=True,
                source="dictionary",
            )
        ],
    )

    BackMatterWriter().rebuild(document, report)

    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name.startswith("Heading 1")
    ]
    assert headings == [
        "Recommended Resources",
        "Index",
        "Glossary",
        "About the Author",
        "Contact",
    ]
    assert "Old index content" not in [
        paragraph.text for paragraph in document.paragraphs
    ]
    assert any(
        element.text == " INDEX "
        for paragraph in document.paragraphs
        for element in paragraph._p.xpath(".//w:instrText")
    )
    assert any(
        paragraph.text == "PowerShell (Technology): A shell."
        for paragraph in document.paragraphs
    )
