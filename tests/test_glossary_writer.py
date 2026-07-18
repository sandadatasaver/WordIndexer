from docx import Document

from wordindexer.glossary import GlossaryBuilder
from wordindexer.glossary_writer import GlossaryWriter
from wordindexer.models import DictionaryEntry


def test_glossary_writer_appends_formatted_word_section(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "glossary.docx"

    source_document = Document()
    source_document.add_paragraph("PowerShell appears here.")
    source_document.save(source)

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            definition="A shell and scripting language.",
            category="Technology",
        )
    ]
    report = GlossaryBuilder().build(source, dictionary)

    output_document = Document()
    output_document.add_paragraph("Main content")
    heading = GlossaryWriter().append(output_document, report)
    output_document.save(output)

    assert heading.text == "Glossary"
    assert heading.style.name == "Heading 1"
    assert output_document.paragraphs[-1].text == (
        "PowerShell (Technology): A shell and scripting language."
    )
    assert output_document.paragraphs[-1].runs[0].bold is True

    reopened = Document(output)
    assert reopened.paragraphs[-2].text == "Glossary"
    assert reopened.paragraphs[-1].text == (
        "PowerShell (Technology): A shell and scripting language."
    )
