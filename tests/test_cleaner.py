from docx import Document

from wordindexer.cleaner import DocumentCleaner


def test_cleaner_removes_named_heading_section_only():
    document = Document()
    document.add_paragraph("Main content")
    document.add_paragraph(
        "Appendix F: Tested With and Tool Requirements",
        style="Heading 1",
    )
    document.add_paragraph("Editorial production note")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Appendix material"
    document.add_paragraph("About the Author", style="Heading 1")
    document.add_paragraph("Biography")

    removed = DocumentCleaner().remove_sections(
        document,
        ["Appendix F: Tested With and Tool Requirements"],
    )

    texts = [paragraph.text for paragraph in document.paragraphs]
    assert removed == ["Appendix F: Tested With and Tool Requirements"]
    assert "Editorial production note" not in texts
    assert "About the Author" in texts
    assert len(document.tables) == 0
