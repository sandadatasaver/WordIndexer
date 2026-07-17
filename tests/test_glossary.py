import json

from docx import Document

from wordindexer.glossary import GlossaryBuilder
from wordindexer.models import DictionaryEntry


def test_glossary_builder_includes_found_and_missing_defined_terms(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("PowerShell is useful.")
    document.save(source)

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            index_as="PowerShell",
            definition="A shell and scripting language.",
            category="Technology",
        ),
        DictionaryEntry(
            term="Pipeline",
            index_as="Pipeline",
            definition="A sequence that passes output between commands.",
            category="Concept",
        ),
        DictionaryEntry(
            term="Ignored",
            index_as="Ignored",
            definition="Should not appear in the glossary.",
            include_in_glossary=False,
        ),
    ]

    report = GlossaryBuilder().build(source, dictionary)

    assert [entry.term for entry in report.entries] == [
        "Pipeline",
        "PowerShell",
    ]
    assert report.entries[0].found is False
    assert report.entries[0].occurrences == 0
    assert report.entries[1].found is True
    assert report.entries[1].occurrences == 1


def test_glossary_report_writes_json_and_csv(tmp_path):
    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("PowerShell")
    document.save(source)

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            definition="A shell.",
        )
    ]
    report = GlossaryBuilder().build(source, dictionary)

    json_path = report.write_json(tmp_path / "glossary.json")
    csv_path = report.write_csv(tmp_path / "glossary.csv")

    payload = json.loads(json_path.read_text(encoding="utf-8"))
    assert payload["entries"][0]["term"] == "PowerShell"
    assert csv_path.read_text(encoding="utf-8").startswith(
        "term,definition,category,aliases,occurrences,found,source"
    )
