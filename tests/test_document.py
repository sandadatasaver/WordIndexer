from docx import Document

from wordindexer.document import DocumentReader


def test_sample_document_exists():
    from pathlib import Path

    assert Path("input/sample.docx").exists()


def test_document_reader():
    reader = DocumentReader("input/sample.docx")

    info = reader.inspect()

    assert info.paragraphs >= 0
    assert info.tables >= 0
    assert info.images >= 0


def test_document_reader_can_traverse_table_cells(tmp_path):
    source = tmp_path / "table.docx"
    document = Document()
    document.add_paragraph("Body paragraph")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("PowerShell in a table")
    document.save(source)

    reader = DocumentReader(source)

    body_only = reader.load_book()
    with_tables = reader.load_book(include_tables=True)

    assert len(body_only.paragraphs) == 1
    assert len(with_tables.paragraphs) == 2

    table_paragraph = with_tables.paragraphs[-1]
    assert table_paragraph.story == "table"
    assert table_paragraph.table_index == 0
    assert table_paragraph.row_index == 0
    assert table_paragraph.cell_index == 0
    assert table_paragraph.cell_paragraph_index == 0
    assert with_tables.paragraph_targets[table_paragraph.index] is not None
