from docx import Document

from wordindexer.index import IndexEngine
from wordindexer.models import DictionaryEntry


def test_index_engine_saves_xe_fields_and_excludes_toc(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "output" / "indexed.docx"

    document = Document()
    document.add_paragraph("Table of Contents", style="Heading 1")
    document.add_paragraph("Chapter 1 ............................ 1")
    document.add_paragraph("PowerShell ........................... 4")
    document.add_paragraph("Chapter 1", style="Heading 1")
    document.add_paragraph(
        "PowerShell appears here. PowerShell appears again."
    )
    document.save(source)

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            index_as="PowerShell",
        )
    ]

    result = IndexEngine().index(source, dictionary, output)

    assert output.exists()
    assert result.toc_detected is True
    assert result.toc_method == "title"
    assert result.body_start == 3
    assert result.terms_found == 1
    assert result.terms_not_found == 0
    assert result.occurrences == 2
    assert result.fields_inserted == 2

    indexed = Document(output)
    assert indexed.paragraphs[4].text == (
        "PowerShell appears here. PowerShell appears again."
    )

    instructions = [
        element.text
        for paragraph in indexed.paragraphs
        for element in paragraph._p.xpath(".//w:instrText")
    ]

    assert instructions == [
        ' XE "',
        "PowerShell",
        '" ',
        ' XE "',
        "PowerShell",
        '" ',
    ]

    toc_instructions = indexed.paragraphs[2]._p.xpath(".//w:instrText")
    assert toc_instructions == []
