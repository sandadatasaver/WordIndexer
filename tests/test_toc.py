from docx import Document

from wordindexer.toc import TOCDetector


def test_toc_detector_finds_body_after_toc_entries():
    document = Document()
    document.add_paragraph("Table of Contents", style="Heading 1")
    document.add_paragraph("Chapter 1 ............................ 1")
    document.add_paragraph("PowerShell ........................... 4")
    document.add_paragraph("Chapter 1", style="Heading 1")
    document.add_paragraph("PowerShell begins here.")

    result = TOCDetector().detect(document)

    assert result.found is True
    assert result.toc_start == 0
    assert result.body_start == 3
    assert result.method == "title"


def test_toc_detector_uses_first_chapter_fallback():
    document = Document()
    document.add_paragraph("Dedication", style="Heading 1")
    document.add_paragraph("Preface", style="Heading 1")
    document.add_paragraph(
        "Chapter 1: Welcome to PowerShell",
        style="Heading 1",
    )
    document.add_paragraph("PowerShell begins here.")

    result = TOCDetector().detect(document)

    assert result.found is False
    assert result.toc_start is None
    assert result.body_start == 2
    assert result.method == "first_chapter_heading"


def test_toc_detector_defaults_to_whole_document_when_missing():
    document = Document()
    document.add_paragraph("PowerShell begins here.")

    result = TOCDetector().detect(document)

    assert result.found is False
    assert result.toc_start is None
    assert result.body_start == 0
    assert result.method == "none"
