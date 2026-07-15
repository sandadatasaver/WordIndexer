from docx import Document

from wordindexer.models import (
    Book,
    DictionaryEntry,
    Heading,
    Paragraph,
)
from wordindexer.search import SearchEngine


def test_search():
    book = Book()

    book.paragraphs.append(
        Paragraph(
            index=0,
            text="PowerShell uses Get-ChildItem.",
            style="Normal",
        )
    )

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            index_as="PowerShell",
        ),
        DictionaryEntry(
            term="Get-ChildItem",
            aliases=["gci"],
            index_as="Get-ChildItem",
        ),
    ]

    engine = SearchEngine(book)
    results = engine.search(dictionary)

    assert len(results["PowerShell"]) == 1
    assert len(results["Get-ChildItem"]) == 1


def test_heading_context():
    book = Book()

    book.paragraphs.append(
        Paragraph(
            index=0,
            text="Chapter One",
            style="Heading 1",
        )
    )

    book.headings.append(
        Heading(
            level=1,
            text="Chapter One",
            paragraph_index=0,
        )
    )

    book.paragraphs.append(
        Paragraph(
            index=1,
            text="PowerShell is powerful.",
            style="Normal",
        )
    )

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            index_as="PowerShell",
        )
    ]

    engine = SearchEngine(book)
    results = engine.search(dictionary)

    assert len(results["PowerShell"]) == 1
    assert results["PowerShell"][0].heading == "Chapter One"


def test_search_returns_every_occurrence():
    book = Book()
    book.paragraphs.append(
        Paragraph(
            index=0,
            text="PowerShell is useful. PowerShell is scriptable.",
        )
    )

    dictionary = [
        DictionaryEntry(term="PowerShell", index_as="PowerShell"),
    ]

    results = SearchEngine(book).search(dictionary)

    assert len(results["PowerShell"]) == 2
    assert [item.matched_text for item in results["PowerShell"]] == [
        "PowerShell",
        "PowerShell",
    ]


def test_search_aliases_use_canonical_index_term():
    book = Book()
    book.paragraphs.append(
        Paragraph(
            index=0,
            text="Use gci first, then use Get-ChildItem.",
        )
    )

    dictionary = [
        DictionaryEntry(
            term="Get-ChildItem",
            aliases=["gci"],
            index_as="Get-ChildItem",
        ),
    ]

    results = SearchEngine(book).search(dictionary)

    assert len(results["Get-ChildItem"]) == 2
    assert {item.term for item in results["Get-ChildItem"]} == {
        "Get-ChildItem",
    }
    assert [item.matched_text for item in results["Get-ChildItem"]] == [
        "gci",
        "Get-ChildItem",
    ]


def test_search_attaches_exact_run_locations():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Power")
    paragraph.add_run("Shell")
    paragraph.add_run(" is useful.")

    book = Book()
    book.paragraphs.append(
        Paragraph(
            index=0,
            text=paragraph.text,
        )
    )

    dictionary = [
        DictionaryEntry(term="PowerShell", index_as="PowerShell"),
    ]

    results = SearchEngine(book, document).search(dictionary)
    match = results["PowerShell"][0]

    assert match.matched_text == "PowerShell"
    assert len(match.locations) == 2
    assert match.locations[0].run_index == 0
    assert match.locations[0].matched_text == "Power"
    assert match.locations[1].run_index == 1
    assert match.locations[1].matched_text == "Shell"


def test_search_resolves_overlapping_terms_globally():
    document = Document()
    paragraph = document.add_paragraph("Windows PowerShell 7 is useful.")

    book = Book()
    book.paragraphs.append(
        Paragraph(
            index=0,
            text=paragraph.text,
        )
    )

    dictionary = [
        DictionaryEntry(
            term="PowerShell",
            index_as="PowerShell",
        ),
        DictionaryEntry(
            term="Windows PowerShell",
            index_as="Windows PowerShell",
        ),
    ]

    results = SearchEngine(book, document).search(dictionary)

    assert len(results["PowerShell"]) == 0
    assert len(results["Windows PowerShell"]) == 1
