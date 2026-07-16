from docx import Document

from wordindexer.index import IndexEngine
from wordindexer.models import DictionaryEntry


def test_index_engine_can_index_table_cell_paragraph(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "indexed.docx"

    document = Document()
    document.add_paragraph("Body text")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run(
        "PowerShell appears inside a table."
    )
    document.save(source)

    result = IndexEngine(
        include_index_field=False,
        include_tables=True,
    ).index(
        source,
        [DictionaryEntry(term="PowerShell", index_as="PowerShell")],
        output,
    )

    assert result.occurrences == 1
    assert result.fields_inserted == 1

    indexed = Document(output)
    cell_paragraph = indexed.tables[0].cell(0, 0).paragraphs[0]
    assert cell_paragraph.text == "PowerShell appears inside a table."
    assert [
        element.text
        for element in cell_paragraph._p.xpath(".//w:instrText")
    ] == [
        ' XE "',
        "PowerShell",
        '" ',
    ]
