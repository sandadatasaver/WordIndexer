from docx import Document

from wordindexer.scanner import RunScanner


def test_run_scanner():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("PowerShell ")
    paragraph.add_run("uses ")
    paragraph.add_run("Get-ChildItem")

    scanner = RunScanner(doc)

    locations = scanner.locate(
        paragraph_index=0,
        term="Get-ChildItem",
    )

    assert len(locations) == 1
    assert locations[0].run_index == 2
    assert locations[0].matched_text == "Get-ChildItem"


def test_scanner_returns_every_occurrence():
    doc = Document()
    paragraph = doc.add_paragraph(
        "PowerShell is useful. PowerShell is scriptable."
    )

    scanner = RunScanner(doc)
    occurrences = scanner.locate_occurrences(0, "PowerShell")

    assert len(occurrences) == 2

    assert [item[0].matched_text for item in occurrences] == [
        "PowerShell",
        "PowerShell",
    ]

    assert [item[0].start for item in occurrences] == [0, 22]


def test_scanner_groups_term_split_across_runs():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Power")
    paragraph.add_run("Shell")
    paragraph.add_run(" automation")

    scanner = RunScanner(doc)
    occurrences = scanner.locate_occurrences(0, "PowerShell")

    assert len(occurrences) == 1
    assert len(occurrences[0]) == 2

    assert occurrences[0][0].run_index == 0
    assert occurrences[0][0].start == 0
    assert occurrences[0][0].end == 5
    assert occurrences[0][0].matched_text == "Power"

    assert occurrences[0][1].run_index == 1
    assert occurrences[0][1].start == 0
    assert occurrences[0][1].end == 5
    assert occurrences[0][1].matched_text == "Shell"