from docx import Document

from wordindexer.index import IndexEngine
from wordindexer.models import DictionaryEntry


def test_table_before_body_boundary_is_not_indexed(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "indexed.docx"

    document = Document()
    document.add_paragraph("Table of Contents", style="Heading 1")

    before_body = document.add_table(rows=1, cols=1)
    before_body.cell(0, 0).text = "PowerShell in the TOC"

    document.add_paragraph("Chapter 1", style="Heading 1")

    after_body = document.add_table(rows=1, cols=1)
    after_body.cell(0, 0).text = "PowerShell in the body"
    document.save(source)

    result = IndexEngine(
        include_index_field=False,
        include_tables=True,
    ).index(
        source,
        [DictionaryEntry(term="PowerShell", index_as="PowerShell")],
        output,
    )

    assert result.body_start == 2
    assert result.occurrences == 1

    indexed = Document(output)
    before_fields = indexed.tables[0].cell(0, 0).paragraphs[0]._p.xpath(
        ".//w:instrText"
    )
    after_fields = indexed.tables[1].cell(0, 0).paragraphs[0]._p.xpath(
        ".//w:instrText"
    )

    assert before_fields == []
    assert [element.text for element in after_fields] == [
        ' XE "',
        "PowerShell",
        '" ',
    ]
