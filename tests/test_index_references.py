from docx import Document

from wordindexer.index import IndexEngine
from wordindexer.models import DictionaryEntry


def test_index_engine_inserts_see_also_once_per_canonical_entry(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "indexed.docx"

    document = Document()
    document.add_paragraph("PowerShell appears. PowerShell appears again.")
    document.save(source)

    entry = DictionaryEntry(
        term="PowerShell",
        index_as="PowerShell",
        see_also=["Windows PowerShell"],
    )

    IndexEngine(include_index_field=False).index(
        source,
        [entry],
        output,
    )

    indexed = Document(output)
    instructions = [
        element.text
        for paragraph in indexed.paragraphs
        for element in paragraph._p.xpath(".//w:instrText")
    ]

    assert instructions.count("PowerShell") == 3
    assert instructions.count("See also Windows PowerShell") == 1
