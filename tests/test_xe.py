from docx import Document

from wordindexer.models import DictionaryEntry, Match, RunLocation
from wordindexer.xe import XEWriter


def test_xe_field_code_uses_canonical_term():
    assert XEWriter.field_code("PowerShell") == 'XE "PowerShell"'
    assert XEWriter.field_code("PowerShell: Basics") == (
        'XE "PowerShell\\: Basics"'
    )


def test_xe_field_code_escapes_special_characters():
    assert XEWriter.field_code('A "quoted"; term') == (
        'XE "A \\\"quoted\\\"\\; term"'
    )


def test_xe_field_code_supports_hierarchy():
    entry = DictionaryEntry(
        term="Get-ChildItem",
        index_as="Get-ChildItem",
        parent="PowerShell",
        subentry="Cmdlets",
    )

    assert XEWriter.hierarchy_text(entry) == (
        "PowerShell:Cmdlets:Get-ChildItem"
    )
    assert XEWriter.field_code_for_entry(entry) == (
        'XE "PowerShell:Cmdlets:Get-ChildItem"'
    )


def test_xe_field_code_supports_see_reference():
    entry = DictionaryEntry(
        term="pwsh",
        index_as="pwsh",
        see="PowerShell",
    )

    assert XEWriter.field_code_for_entry(entry) == (
        'XE "pwsh" \\t "See PowerShell"'
    )
    assert XEWriter.instruction_parts_for_entry(entry) == [
        ' XE "',
        "pwsh",
        '" \\t "',
        "See PowerShell",
        '" ',
    ]


def test_xe_writer_inserts_field_at_match_location():
    document = Document()
    paragraph = document.add_paragraph("Before PowerShell after")

    match = Match(
        term="PowerShell",
        matched_text="PowerShell",
        paragraph_index=0,
        paragraph_text=paragraph.text,
        locations=[
            RunLocation(
                paragraph_index=0,
                run_index=0,
                start=7,
                end=17,
                matched_text="PowerShell",
            )
        ],
    )

    XEWriter().insert_match(paragraph, match)

    assert paragraph.text == "Before PowerShell after"
    instructions = paragraph._p.xpath(".//w:instrText")
    assert [instruction.text for instruction in instructions] == [
        ' XE "',
        "PowerShell",
        '" ',
    ]


def test_xe_field_code_supports_see_also_reference():
    entry = DictionaryEntry(
        term="PowerShell",
        index_as="PowerShell",
        see_also=["Windows PowerShell", "PowerShell Core"],
    )

    assert XEWriter.see_also_text(entry) == (
        "See also Windows PowerShell; PowerShell Core"
    )
    assert XEWriter.field_code_for_see_also(entry) == (
        'XE "PowerShell" \\t "See also Windows PowerShell; PowerShell Core"'
    )
    assert XEWriter.instruction_parts_for_see_also(entry) == [
        ' XE "',
        "PowerShell",
        '" \\t "',
        "See also Windows PowerShell; PowerShell Core",
        '" ',
    ]


def test_xe_writer_inserts_see_reference():
    document = Document()
    paragraph = document.add_paragraph("Use pwsh here")

    entry = DictionaryEntry(
        term="pwsh",
        index_as="pwsh",
        see="PowerShell",
    )
    match = Match(
        term="pwsh",
        matched_text="pwsh",
        paragraph_index=0,
        paragraph_text=paragraph.text,
        dictionary_entry=entry,
        locations=[
            RunLocation(
                paragraph_index=0,
                run_index=0,
                start=4,
                end=8,
                matched_text="pwsh",
            )
        ],
    )

    XEWriter().insert_match(paragraph, match)

    assert [element.text for element in paragraph._p.xpath(".//w:instrText")] == [
        ' XE "',
        "pwsh",
        '" \\t "',
        "See PowerShell",
        '" ',
    ]


def test_xe_writer_inserts_see_also_when_requested():
    document = Document()
    paragraph = document.add_paragraph("PowerShell is useful")

    entry = DictionaryEntry(
        term="PowerShell",
        index_as="PowerShell",
        see_also=["Windows PowerShell"],
    )
    match = Match(
        term="PowerShell",
        matched_text="PowerShell",
        paragraph_index=0,
        paragraph_text=paragraph.text,
        dictionary_entry=entry,
        locations=[
            RunLocation(
                paragraph_index=0,
                run_index=0,
                start=0,
                end=10,
                matched_text="PowerShell",
            )
        ],
    )

    XEWriter().insert_match(
        paragraph,
        match,
        include_see_also=True,
    )

    assert [element.text for element in paragraph._p.xpath(".//w:instrText")] == [
        ' XE "',
        "PowerShell",
        '" ',
        ' XE "',
        "PowerShell",
        '" \\t "',
        "See also Windows PowerShell",
        '" ',
    ]


def test_xe_writer_indexes_alias_under_canonical_term():
    document = Document()
    paragraph = document.add_paragraph("Use gci here")

    match = Match(
        term="Get-ChildItem",
        matched_text="gci",
        paragraph_index=0,
        paragraph_text=paragraph.text,
        locations=[
            RunLocation(
                paragraph_index=0,
                run_index=0,
                start=4,
                end=7,
                matched_text="gci",
            )
        ],
    )

    XEWriter().insert_match(paragraph, match)

    instructions = paragraph._p.xpath(".//w:instrText")
    assert [instruction.text for instruction in instructions] == [
        ' XE "',
        "Get-ChildItem",
        '" ',
    ]
