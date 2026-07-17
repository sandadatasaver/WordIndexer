from docx import Document

from wordindexer.index_field import IndexFieldWriter


def test_index_field_writer_appends_pending_index_field(tmp_path):
    document = Document()
    document.add_paragraph("Indexed manuscript content")

    paragraph = IndexFieldWriter().insert_index_field(document)

    assert document.paragraphs[-2].text == "Index"
    assert document.paragraphs[-2].style.name == "Heading 1"
    assert document.paragraphs[-2].paragraph_format.page_break_before is True
    assert paragraph.text == ""
    assert paragraph.paragraph_format.page_break_before is None

    instructions = paragraph._p.xpath(".//w:instrText")
    assert len(instructions) == 1
    assert instructions[0].text == " INDEX "

    field_types = [
        element.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType"
        )
        for element in paragraph._p.xpath(".//w:fldChar")
    ]
    assert field_types == ["begin", "separate", "end"]

    output = tmp_path / "index_field.docx"
    document.save(output)

    reopened = Document(output)
    reopened_instructions = [
        element.text
        for p in reopened.paragraphs
        for element in p._p.xpath(".//w:instrText")
    ]
    assert reopened_instructions == [" INDEX "]
