import json

from docx import Document

from wordindexer.models import DictionaryEntry
from wordindexer.reports import AnalysisReport, ReportBuilder


def test_report_builder_is_non_modifying_and_counts_coverage(tmp_path):
    source = tmp_path / "analysis.docx"
    document = Document()
    document.add_paragraph("Table of Contents", style="Heading 1")
    document.add_paragraph("Chapter 1 ............................ 1")
    document.add_paragraph("Chapter 1", style="Heading 1")
    document.add_paragraph("Windows PowerShell is useful.")
    document.save(source)

    dictionary = [
        DictionaryEntry(term="PowerShell", index_as="PowerShell"),
        DictionaryEntry(
            term="Windows PowerShell",
            index_as="Windows PowerShell",
        ),
        DictionaryEntry(term="Missing", index_as="Missing"),
    ]

    report = ReportBuilder().build(source, dictionary)

    assert report.body_start == 2
    assert report.toc_detected is True
    assert report.toc_method == "title"
    assert report.ignored_paragraphs == 2
    assert report.dictionary_entries == 3
    assert report.terms_found == 1
    assert report.terms_missing == 2
    assert report.total_occurrences == 1
    assert report.overlaps_resolved == 1
    assert report.term_counts["Windows PowerShell"] == 1
    assert report.missing_terms == ["PowerShell", "Missing"]
    assert source.exists()


def test_report_can_be_written_as_json(tmp_path):
    report = AnalysisReport(
        input_path="source.docx",
        total_paragraphs=10,
        body_start=3,
        ignored_paragraphs=3,
        toc_detected=True,
        toc_method="title",
        dictionary_entries=2,
        terms_found=1,
        terms_missing=1,
        total_occurrences=4,
        overlaps_resolved=1,
        term_counts={"PowerShell": 4},
        missing_terms=["Missing"],
    )

    output = report.write_json(tmp_path / "report.json")
    payload = json.loads(output.read_text(encoding="utf-8"))

    assert payload["total_occurrences"] == 4
    assert payload["missing_terms"] == ["Missing"]
